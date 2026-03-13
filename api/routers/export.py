"""Document export API router — preview and download."""

from __future__ import annotations

import io
import logging

from fastapi import APIRouter
from fastapi.responses import StreamingResponse

from api.models import ExportRequest

logger = logging.getLogger(__name__)
router = APIRouter()


@router.post("/preview")
async def preview_document(req: ExportRequest):
    """Render markdown content to HTML for preview."""
    try:
        import markdown
        html = markdown.markdown(
            req.content,
            extensions=["tables", "fenced_code", "codehilite", "toc"],
        )
        styled_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: -apple-system, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background: #f5f5f5; }}
pre {{ background: #f5f5f5; padding: 12px; border-radius: 4px; overflow-x: auto; }}
code {{ background: #f0f0f0; padding: 2px 4px; border-radius: 3px; }}
h1,h2,h3 {{ color: #333; }}
</style></head><body>{html}</body></html>"""
        return {"ok": True, "data": {"html": styled_html}}
    except ImportError:
        # Fallback: return raw markdown wrapped in <pre>
        return {"ok": True, "data": {"html": f"<pre>{req.content}</pre>"}}
    except Exception as e:
        return {"ok": False, "message": str(e)}


@router.post("/download")
async def download_document(req: ExportRequest):
    """Generate downloadable file (markdown or docx)."""
    if req.format == "markdown":
        content_bytes = req.content.encode("utf-8")
        filename = f"{req.filename}.md"
        media_type = "text/markdown"
    elif req.format == "docx":
        try:
            from docx import Document
            doc = Document()
            for line in req.content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped:
                    doc.add_paragraph(stripped)
            buf = io.BytesIO()
            doc.save(buf)
            content_bytes = buf.getvalue()
            filename = f"{req.filename}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except ImportError:
            return {"ok": False, "message": "python-docx 未安装，无法导出 Word 格式"}
    else:
        return {"ok": False, "message": f"不支持的格式: {req.format}"}

    return StreamingResponse(
        io.BytesIO(content_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
