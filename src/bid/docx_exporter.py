"""DOCX Exporter - 导出商务文件为 Word 文档.

Uses python-docx to generate professional Word documents.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def export_to_docx(
    outline: List[Dict[str, Any]],
    content: Dict[str, str],
    attachments: Optional[Dict[str, str]] = None,
    output_path: str = "商务文件.docx",
    project_name: str = "",
    project_code: str = "",
) -> str:
    """导出商务文件为 Word 文档.
    
    Args:
        outline: 大纲列表
        content: 章节内容映射 {section_id: content}
        attachments: 附件映射 {section_id: file_path}
        output_path: 输出文件路径
        project_name: 项目名称
        project_code: 项目编号
    
    Returns:
        输出文件路径
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Inches, Pt
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise ImportError("python-docx is required for DOCX export")
    
    attachments = attachments or {}
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
    
    title_text = f"{project_name} 商务文件" if project_name else "商务文件"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if project_code:
        subtitle = doc.add_paragraph(f"项目编号：{project_code}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for item in outline:
        section_id = item.get("id", "")
        section_title = item.get("title", "")
        level = item.get("level", 1)
        
        if level == 1:
            doc.add_heading(section_title, level=1)
        elif level == 2:
            doc.add_heading(section_title, level=2)
        else:
            doc.add_heading(section_title, level=3)
        
        section_content = content.get(section_id, "")
        if section_content:
            for para_text in section_content.split("\n\n"):
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    para.paragraph_format.line_spacing = 1.5
        
        if section_id in attachments:
            attachment_path = attachments[section_id]
            if Path(attachment_path).exists():
                suffix = Path(attachment_path).suffix.lower()
                if suffix in (".jpg", ".jpeg", ".png", ".bmp", ".gif"):
                    try:
                        doc.add_paragraph()
                        doc.add_picture(attachment_path, width=Inches(5.5))
                        caption = doc.add_paragraph(f"图：{section_title}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        logger.warning(f"Failed to add image {attachment_path}: {e}")
                        doc.add_paragraph(f"[附件：{Path(attachment_path).name}]")
                else:
                    doc.add_paragraph(f"[附件：{Path(attachment_path).name}]")
    
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)
    logger.info(f"Document exported to: {output_path}")
    return output_path


def export_outline_only(
    outline: List[Dict[str, Any]],
    output_path: str = "商务文件大纲.docx",
    project_name: str = "",
) -> str:
    """仅导出大纲（不含内容）.
    
    Args:
        outline: 大纲列表
        output_path: 输出文件路径
        project_name: 项目名称
    
    Returns:
        输出文件路径
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")
    
    doc = Document()
    
    title_text = f"{project_name} 商务文件大纲" if project_name else "商务文件大纲"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for item in outline:
        section_id = item.get("id", "")
        section_title = item.get("title", "")
        level = item.get("level", 1)
        
        indent = "  " * (level - 1)
        text = f"{indent}{section_id} {section_title}"
        
        para = doc.add_paragraph(text)
        para.paragraph_format.left_indent = Cm(0.5 * (level - 1))
    
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)
    logger.info(f"Outline exported to: {output_path}")
    return output_path


def create_empty_template() -> str:
    """创建空白商务文件模板.
    
    Returns:
        模板内容（用于 BidTemplate）
    """
    return """{{公司名称}}

关于{{项目名称}}项目的承诺书

致：{{招标单位}}

我公司郑重承诺：

一、我公司具备完成本项目所需的资质、能力和经验。

二、我公司提供的所有投标文件及资料均真实、有效、合法。

三、我公司将严格按照招标文件要求履行合同义务。

四、如有违反上述承诺，我公司愿承担相应的法律责任。

特此承诺。

承诺单位：{{公司名称}}（盖章）
法定代表人或授权代表：____________（签字）
日期：____年____月____日"""
