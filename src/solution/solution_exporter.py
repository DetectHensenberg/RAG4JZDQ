"""Solution Exporter — 将方案导出为 Word 文档.

使用 python-docx 生成格式化的 Word 文件，
包含封面、目录结构、各章节内容。
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


def export_to_docx(
    outline: List[Dict[str, Any]],
    content: Dict[str, str],
    output_path: str,
    project_name: str = "",
    project_type: str = "",
) -> str:
    """导出方案为 Word 文档.

    Args:
        outline: 方案大纲
        content: 各章节内容 {section_id: markdown}
        output_path: 输出文件路径
        project_name: 项目名称（用于封面）
        project_type: 项目类型

    Returns:
        输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ValueError("需要安装 python-docx: pip install python-docx")

    doc = Document()

    # ── 设置默认样式 ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # ── 封面页 ────────────────────────────────────────────────────
    _add_cover_page(doc, project_name, project_type)
    doc.add_page_break()

    # ── 目录页（文本模拟） ────────────────────────────────────────
    doc.add_heading("目  录", level=1)
    for item in outline:
        level = item.get("level", 1)
        title = item.get("title", "")
        indent = "    " * (level - 1)
        p = doc.add_paragraph(f"{indent}{title}")
        p.paragraph_format.space_after = Pt(2)
        if level == 1:
            for run in p.runs:
                run.bold = True
    doc.add_page_break()

    # ── 正文内容 ──────────────────────────────────────────────────
    for item in outline:
        section_id = item.get("id", "")
        title = item.get("title", "")
        level = min(item.get("level", 1), 3)

        # 添加标题
        doc.add_heading(title, level=level)

        # 添加内容
        section_content = content.get(section_id, "")
        if section_content:
            _add_markdown_content(doc, section_content)

    # ── 保存 ─────────────────────────────────────────────────────
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"方案导出成功: {output_path}")
    return output_path


def _add_cover_page(
    doc: Any,
    project_name: str,
    project_type: str,
) -> None:
    """添加封面页.

    Args:
        doc: Document 对象
        project_name: 项目名称
        project_type: 项目类型
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 空行留白
    for _ in range(6):
        doc.add_paragraph("")

    # 项目名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name or "技术方案")
    run.bold = True
    run.font.size = Pt(26)

    doc.add_paragraph("")

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("技术解决方案")
    run.font.size = Pt(18)

    if project_type:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"项目类型: {project_type}")
        run.font.size = Pt(14)

    # 日期
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(14)


def _add_markdown_content(doc: Any, md_text: str) -> None:
    """将 Markdown 文本转换为 Word 段落.

    支持基础 Markdown 格式：
    - 段落
    - 列表（- / * / 1.）
    - 粗体（**text**）
    - 代码块（```）

    Args:
        doc: Document 对象
        md_text: Markdown 格式文本
    """
    from docx.shared import Pt

    lines = md_text.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        stripped = line.strip()

        # 代码块处理
        if stripped.startswith("```"):
            if in_code_block:
                # 结束代码块
                if code_buffer:
                    code_text = "\n".join(code_buffer)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.size = Pt(9)
                    run.font.name = "Consolas"
                    code_buffer.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 空行
        if not stripped:
            continue

        # 列表
        if re.match(r"^[-*]\s", stripped):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)


def _add_formatted_text(paragraph: Any, text: str) -> None:
    """添加支持粗体标记的格式化文本.

    Args:
        paragraph: Word 段落对象
        text: 可能包含 **bold** 标记的文本
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
