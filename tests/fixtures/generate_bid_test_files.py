"""Generate mock DOCX and Excel files for bid assistant testing.

Run this script to create test fixtures:
    python tests/fixtures/generate_bid_test_files.py
"""

from pathlib import Path

OUTPUT_DIR = Path(__file__).parent / "bid_test_data"
OUTPUT_DIR.mkdir(exist_ok=True)


def generate_tender_docx() -> Path:
    """Generate a mock tender (招标) document with disqualification items."""
    from docx import Document

    doc = Document()
    doc.add_heading("XX市智慧城市集成项目招标文件", level=1)

    doc.add_heading("第一章 投标人须知", level=2)
    doc.add_paragraph(
        "一、投标人应具备以下资格条件：\n"
        "1. 投标人须具有独立法人资格，并提供有效的营业执照副本。\n"
        "2. 投标人须具有信息系统集成及服务资质证书（三级及以上）。\n"
        "3. 投标人近三年内无重大违法违规记录，提供承诺书。"
    )

    doc.add_heading("第二章 符合性审查", level=2)
    doc.add_paragraph(
        "符合性审查内容如下，不满足任一条即作废标处理：\n"
        "1. 投标文件须按规定格式编制，使用A4纸装订成册。\n"
        "2. 投标文件须加盖投标人公章及法定代表人签字或盖章。\n"
        "3. 投标保证金须在投标截止时间前到账，否则废标。\n"
        "4. 投标有效期不少于90天。"
    )

    doc.add_heading("第三章 资格性审查", level=2)
    doc.add_paragraph(
        "资格性审查要求如下：\n"
        "1. 投标人须提供近三年经审计的财务报表。\n"
        "2. 投标人须提供社保缴纳证明。\n"
        "3. 投标人应具备ISO9001质量管理体系认证。\n"
        "未通过资格性审查的投标文件，按无效投标处理。"
    )

    doc.add_heading("第四章 技术要求（实质性要求）", level=2)
    doc.add_paragraph(
        "★ 系统应支持不少于1000个终端设备同时在线接入。\n"
        "★ 系统响应时间不超过3秒。\n"
        "★ 数据存储容量不少于100TB，支持横向扩展。\n"
        "3. 系统应支持7×24小时不间断运行。\n"
        "4. 提供三年免费质保服务，质保期内免费上门维修。"
    )

    doc.add_heading("第五章 商务要求", level=2)
    doc.add_paragraph(
        "1. 项目实施周期不超过180天。\n"
        "2. 投标人须提供不少于3个同类项目业绩证明。\n"
        "★ 投标人须承诺项目经理具备PMP或同等资质认证。\n"
        "3. 付款方式：合同签订后支付30%，验收合格后支付60%，质保期满后支付10%。"
    )

    path = OUTPUT_DIR / "mock_tender.docx"
    doc.save(str(path))
    print(f"Generated: {path}")
    return path


def generate_bid_docx() -> Path:
    """Generate a mock bid (投标) response document."""
    from docx import Document

    doc = Document()
    doc.add_heading("XX市智慧城市集成项目投标文件", level=1)
    doc.add_paragraph("投标人：九洲科技股份有限公司")

    doc.add_heading("第一部分 资格证明材料", level=2)
    doc.add_paragraph(
        "一、营业执照\n"
        "九洲科技股份有限公司，统一社会信用代码：91510100XXXXXXXXXX，"
        "营业执照有效期至2030年12月31日。\n\n"
        "二、信息系统集成及服务资质\n"
        "我公司持有信息系统集成及服务资质证书（二级），证书编号：XZ-2023-XXXX。\n\n"
        "三、无违法违规承诺书\n"
        "我公司承诺近三年内无重大违法违规记录。"
    )

    doc.add_heading("第二部分 商务响应", level=2)
    doc.add_paragraph(
        "一、投标保证金\n"
        "已于2024年12月1日通过银行转账缴纳投标保证金。\n\n"
        "二、投标有效期\n"
        "本投标文件自投标截止日起90天内有效。\n\n"
        "三、项目实施周期\n"
        "我公司承诺在合同签订后150天内完成项目实施。\n\n"
        "四、项目经理资质\n"
        "拟派项目经理张工，持有PMP认证证书。\n\n"
        "五、同类项目业绩\n"
        "1. 成都市智慧交通项目（2022年，合同金额800万元）\n"
        "2. 绵阳市智慧园区项目（2023年，合同金额650万元）\n"
        "3. 德阳市数据中心项目（2023年，合同金额520万元）"
    )

    doc.add_heading("第三部分 技术方案", level=2)
    doc.add_paragraph(
        "一、系统架构\n"
        "采用微服务分布式架构，支持2000个终端设备同时在线接入。\n\n"
        "二、系统性能\n"
        "系统响应时间不超过2秒，满足招标要求。\n\n"
        "三、数据存储\n"
        "采用分布式存储方案，初始容量120TB，支持PB级横向扩展。\n\n"
        "四、系统可靠性\n"
        "系统支持7×24小时不间断运行，可用性达99.99%。\n\n"
        "五、质保服务\n"
        "提供三年免费质保服务，质保期内4小时响应、24小时上门。"
    )

    # Intentionally missing: ISO9001 certification, financial reports, social security proof
    # This creates testable gaps for the review engine

    path = OUTPUT_DIR / "mock_bid.docx"
    doc.save(str(path))
    print(f"Generated: {path}")
    return path


def generate_achievement_excel() -> Path:
    """Generate a mock Excel file with achievement records for batch import."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "业绩数据"

    headers = ["项目名称", "项目内容", "项目金额", "合同签订时间", "验收时间", "甲方联系人", "联系方式", "标签"]
    ws.append(headers)

    records = [
        ["成都市智慧交通管理平台", "智慧交通信号控制及监控系统集成", 800, "2022-03-15", "2022-12-20", "李明", "13800001111", "智慧交通,集成"],
        ["绵阳市智慧园区综合管理系统", "园区安防、能耗、停车等子系统集成", 650, "2023-01-10", "2023-09-30", "王强", "13800002222", "智慧园区,安防"],
        ["德阳市政务数据中心建设", "数据中心基础设施及云平台建设", 520, "2023-06-01", "2024-03-15", "赵丽", "13800003333", "数据中心,云计算"],
        ["遂宁市智慧城管平台", "城市管理综合信息平台开发与集成", 380, "2022-09-01", "2023-06-30", "刘伟", "13800004444", "智慧城市,软件开发"],
        ["泸州市应急指挥系统", "应急指挥调度通信系统建设", 450, "2023-04-20", "2024-01-15", "陈静", "13800005555", "应急指挥,通信"],
        ["宜宾市智慧安防项目", "城市视频监控及智能分析系统", 920, "2022-11-01", "2023-08-20", "张华", "13800006666", "安防,人工智能"],
        ["乐山市教育信息化项目", "教育城域网及智慧校园平台", 280, "2023-08-15", "2024-06-30", "周敏", "13800007777", "教育,信息化"],
    ]

    for row in records:
        ws.append(row)

    path = OUTPUT_DIR / "mock_achievements.xlsx"
    wb.save(str(path))
    print(f"Generated: {path}")
    return path


def generate_achievement_csv() -> Path:
    """Generate a mock CSV file with achievement records."""
    import csv

    path = OUTPUT_DIR / "mock_achievements.csv"
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["项目名称", "项目内容", "合同金额", "签订时间", "验收时间", "联系人", "电话", "标签"])
        writer.writerow(["南充市智慧水务平台", "智慧水务监测及管理系统建设", "350", "2023-05-10", "2024-02-28", "吴刚", "13900001111", "智慧水务"])
        writer.writerow(["广元市数字政务平台", "政务服务一体化平台建设", "420", "2023-07-01", "2024-04-15", "郑芳", "13900002222", "数字政务,软件开发"])

    print(f"Generated: {path}")
    return path


if __name__ == "__main__":
    generate_tender_docx()
    generate_bid_docx()
    generate_achievement_excel()
    generate_achievement_csv()
    print("All test fixtures generated successfully.")
